# src/multi_agentic_graph_rag/infrastructure/documents/docx_parser.py

from __future__ import annotations

from pathlib import Path

from docx import Document

from multi_agentic_graph_rag.domain.documents import ParsedBlock, ParsedDocument
from multi_agentic_graph_rag.infrastructure.documents.normalization import (
    normalize_text,
    sha256_file,
)


class DocxParser:
    parser_name = "python-docx"
    parser_version = "1.2.0"
    supported_extensions = frozenset({".docx"})

    def parse(self, path: Path) -> ParsedDocument:
        path = path.resolve()
        checksum = sha256_file(path)
        document = Document(str(path))

        blocks: list[ParsedBlock] = []
        section_path: list[str] = []
        cursor = 0

        for index, paragraph in enumerate(document.paragraphs, start=1):
            raw_text = paragraph.text or ""
            normalized = normalize_text(raw_text)

            if not normalized:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""

            if style_name.lower().startswith("heading"):
                level = _heading_level(style_name)
                if level is not None:
                    section_path = [*section_path[: level - 1], normalized]

            start = cursor
            end = start + len(raw_text)
            cursor = end + 1

            blocks.append(
                ParsedBlock(
                    source_path=str(path),
                    source_checksum=checksum,
                    page_number=None,
                    section_path=tuple(section_path),
                    paragraph_number=index,
                    character_start=start,
                    character_end=end,
                    raw_text=raw_text,
                    normalized_text=normalized,
                    parser_name=self.parser_name,
                    parser_version=self.parser_version,
                    metadata={"style_name": style_name},
                )
            )

        return ParsedDocument(
            source_path=str(path),
            source_checksum=checksum,
            parser_name=self.parser_name,
            parser_version=self.parser_version,
            blocks=tuple(blocks),
        )


def _heading_level(style_name: str) -> int | None:
    parts = style_name.split()
    if len(parts) >= 2 and parts[-1].isdigit():
        return max(1, int(parts[-1]))
    return 1
